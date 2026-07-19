"""§9 export: the drafted document as DOCX (advocates live in Word) or markdown.

DOCX layout: title, Heading-2 sections with markdown-lite rendering, per-section
citation endnotes, a verification summary table, and the disclaimer.
# ponytail: python-docx has no footnote API — citations render as small-type
# numbered endnote paragraphs per section, visually footnotes.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MARKER = re.compile(r"(\[\d+\])")

_BADGE = {"quote_verified": "✓ quote verified", "entailed": "◐ entailed",
          "unverified": "⚠ unverified", None: "", "": ""}


def _add_runs(paragraph, text: str) -> None:
    """Render **bold** spans and superscript [n] markers into runs."""
    for chunk in _MARKER.split(text):
        if _MARKER.fullmatch(chunk):
            paragraph.add_run(chunk).font.superscript = True
            continue
        for i, part in enumerate(_BOLD.split(chunk)):
            if part:
                paragraph.add_run(part).bold = i % 2 == 1


def _add_markdown_lite(doc: Document, text: str) -> None:
    for block in (b.strip() for b in text.split("\n")):
        if not block:
            continue
        if block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith(("- ", "* ")):
            _add_runs(doc.add_paragraph(style="List Bullet"), block[2:])
        else:
            _add_runs(doc.add_paragraph(), block)


def _verification_counts(sections: list[dict]) -> tuple[int, int, int]:
    total = verified = 0
    for s in sections:
        for c in s.get("citations", []):
            total += 1
            verified += c.get("verified") in ("quote_verified", "entailed")
    flagged = sum(1 for s in sections if s.get("needs_review"))
    return total, verified, flagged


def job_to_docx(job: dict, sections: list[dict], disclaimer: str) -> bytes:
    doc = Document()
    doc.add_heading(job.get("prompt", "Drafted document")[:120], level=1)

    for sec in sections:
        doc.add_heading(sec["title"], level=2)
        _add_markdown_lite(doc, sec.get("text", ""))
        for i, c in enumerate(sec.get("citations", []), start=1):
            note = doc.add_paragraph()
            run = note.add_run(
                f"{i}. {c.get('source_file') or 'unknown source'} — “{c.get('quote', '')}”"
                + (f"  [{_BADGE.get(c.get('verified'), '')}]" if c.get("verified") else ""))
            run.font.size = Pt(8)
            run.italic = True

    total, verified, flagged = _verification_counts(sections)
    doc.add_heading("Verification", level=2)
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Total citations"
    table.rows[0].cells[1].text = "Verified"
    table.rows[0].cells[2].text = "Flagged sections"
    table.rows[1].cells[0].text = str(total)
    table.rows[1].cells[1].text = f"{verified} ({round(100 * verified / total) if total else 0}%)"
    table.rows[1].cells[2].text = str(flagged)

    doc.add_paragraph().add_run(disclaimer).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def job_to_markdown(job: dict, sections: list[dict], disclaimer: str) -> str:
    parts = [f"# {job.get('prompt', 'Drafted document')[:120]}"]
    for sec in sections:
        cites = "\n".join(
            f"{i}. _{c.get('source_file') or 'unknown source'}_ — \"{c.get('quote', '')}\""
            for i, c in enumerate(sec.get("citations", []), start=1))
        flag = " ⚠ needs review" if sec.get("needs_review") else ""
        parts.append(f"## {sec['title']}{flag}\n\n{sec.get('text', '')}"
                     + (f"\n\n{cites}" if cites else ""))
    total, verified, flagged = _verification_counts(sections)
    parts.append(f"---\n\n**Verification:** {total} citations · {verified} verified · "
                 f"{flagged} flagged sections\n\n_{disclaimer}_")
    return "\n\n".join(parts)
