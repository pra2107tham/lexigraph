"""§9 export: DOCX round-trip (headings, endnotes, verification table,
disclaimer) and the markdown snapshot."""

import io

from docx import Document

JOB = {"_id": "job-1", "prompt": "Summarize the appointment letter", "status": "done"}
SECTIONS = [
    {"section_id": "s1", "title": "Salary", "instructions": "",
     "text": "Pay is **12 LPA** [1].\n### Sub-point\n- monthly instalments [1]",
     "citations": [{"parent_id": "p1", "quote": "12,00,000 per annum",
                    "source_file": "letter.pdf", "verified": "quote_verified"}],
     "needs_review": False},
    {"section_id": "s2", "title": "Notice", "instructions": "",
     "text": "Sixty days [1].",
     "citations": [{"parent_id": "p1", "quote": "sixty (60) days",
                    "source_file": "letter.pdf", "verified": "unverified"}],
     "needs_review": True},
]
DISCLAIMER = "AI-assisted draft — verify before use; not legal advice."


def test_docx_roundtrip():
    from app.export.docx_export import job_to_docx

    doc = Document(io.BytesIO(job_to_docx(JOB, SECTIONS, DISCLAIMER)))
    texts = [p.text for p in doc.paragraphs]
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    assert "Summarize the appointment letter" in headings[0]
    assert "Salary" in headings and "Notice" in headings and "Sub-point" in headings
    assert any("letter.pdf" in t and "12,00,000" in t for t in texts)  # endnote
    assert texts[-1] == DISCLAIMER
    table = doc.tables[0]
    assert table.rows[1].cells[0].text == "2"        # total citations
    assert table.rows[1].cells[1].text.startswith("1 (50%)")
    assert table.rows[1].cells[2].text == "1"        # flagged sections


def test_markdown_export():
    from app.export.docx_export import job_to_markdown

    md = job_to_markdown(JOB, SECTIONS, DISCLAIMER)
    assert md.startswith("# Summarize the appointment letter")
    assert "## Notice ⚠ needs review" in md
    assert '1. _letter.pdf_ — "sixty (60) days"' in md
    assert "2 citations · 1 verified · 1 flagged sections" in md
    assert md.rstrip().endswith(f"_{DISCLAIMER}_")


def test_export_route(client, fake_db):
    from app.stores import mongo

    mongo.jobs().insert_one({**JOB, "document": "x",
                             "outline": {"approved": True, "sections": [
                                 {"section_id": s["section_id"], "title": s["title"]}
                                 for s in SECTIONS]}})
    for s in SECTIONS:
        mongo.drafted_sections().insert_one({**s, "job_id": "job-1"})

    docx = client.get("/jobs/job-1/export?format=docx")
    assert docx.status_code == 200
    assert docx.headers["content-disposition"].endswith('.docx"')
    assert docx.content[:2] == b"PK"  # zip container

    md = client.get("/jobs/job-1/export?format=md")
    assert md.status_code == 200 and "## Salary" in md.text

    assert client.get("/jobs/job-1/export?format=pdf").status_code == 422

    mongo.jobs().update_one({"_id": "job-1"}, {"$set": {"status": "running"}})
    assert client.get("/jobs/job-1/export").status_code == 409
